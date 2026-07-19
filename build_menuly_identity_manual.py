from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\julic\Documents\carta")
LOGO_PATH = Path(r"C:\Users\julic\Downloads\WhatsApp Image 2026-07-19 at 12.45.51.jpeg")
OUT_PATH = ROOT / "Menuly_Brand_Book.docx"


NAVY = "0C1F30"
ORANGE = "F0643A"
IVORY = "FBF8F3"
GRAFITO = "374151"
SLATE = "64748B"
LIGHT = "E7E2DA"


def fmt_run(run, size=11, bold=False, color="111827", font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ("Heading 1", 18, NAVY, 14, 6),
    ("Heading 2", 13, NAVY, 10, 4),
]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)


def add_page_break():
    doc.add_page_break()


def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        r = p.add_run()
        r.add_picture(str(LOGO_PATH), width=Inches(5.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("MENULY")
    fmt_run(r, size=10, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BRAND BOOK")
    fmt_run(r, size=24, bold=False, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Identity system for interactive digital restaurant menus")
    fmt_run(r, size=11, color=SLATE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Modern. Premium. Clear.")
    fmt_run(r, size=11, bold=True, color=GRAFITO)


def section_title(code, title, subtitle):
    p = doc.add_paragraph()
    r = p.add_run(code.upper())
    fmt_run(r, size=8.5, bold=True, color=ORANGE)
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(title)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    fmt_run(r, size=10.5, color=SLATE)


def bullets(items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        r = p.add_run("• ")
        fmt_run(r, size=10.5, color=ORANGE)
        r = p.add_run(item)
        fmt_run(r, size=10.5, color="111827")


add_cover()
add_page_break()

section_title("01", "Brand essence", "What Menuly should feel like at first glance")
bullets([
    "A premium SaaS brand for restaurants that combines technology, elegance and speed.",
    "The identity should feel calm, modern and trustworthy, never loud or playful.",
    "The logo is the anchor: simple, memorable and easy to apply everywhere.",
])

add_page_break()
section_title("02", "Logo system", "Primary lockup, symbol and behavior")
bullets([
    "Primary logo: horizontal lockup with symbol + wordmark.",
    "Secondary logo: symbol only for favicon, app icon and compact placements.",
    "Preferred background: white, ivory or very dark navy.",
    "Clear space: keep at least the height of the M around the logo on every side.",
    "Do not stretch, recolor or add effects.",
])

p = doc.add_paragraph()
p.style = "Heading 2"
p.add_run("Logo rules")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.autofit = False
for i, h in enumerate(["Do", "Dont"]):
    c = tbl.rows[0].cells[i]
    c.text = h
    shade(c, "F3F4F6")
    margins(c)
for cell in tbl.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        fmt_run(run, size=10, bold=True, color=GRAFITO)
do_cell, dont_cell = tbl.add_row().cells
do_cell.text = "Use the horizontal mark on website headers, decks and emails.\nUse the symbol alone in small spaces.\nKeep generous breathing room."
dont_cell.text = "Do not rotate.\nDo not add shadows, bevels or gradients.\nDo not place on busy photos."
for c in [do_cell, dont_cell]:
    margins(c)
    for p in c.paragraphs:
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            fmt_run(run, size=10, color="111827")

add_page_break()
section_title("03", "Color", "A restrained palette with one warm accent")
color_table = doc.add_table(rows=1, cols=4)
color_table.style = "Table Grid"
color_table.autofit = False
for i, h in enumerate(["Token", "Role", "HEX", "Use"]):
    c = color_table.rows[0].cells[i]
    c.text = h
    shade(c, "F3F4F6")
    margins(c)
for run in color_table.rows[0].cells[0].paragraphs[0].runs + color_table.rows[0].cells[1].paragraphs[0].runs + color_table.rows[0].cells[2].paragraphs[0].runs + color_table.rows[0].cells[3].paragraphs[0].runs:
    fmt_run(run, size=10, bold=True, color=GRAFITO)

palette = [
    ("Navy", "Primary ink", f"#{NAVY}", "Wordmark, headings, dark UI"),
    ("Coral", "Accent", f"#{ORANGE}", "Details, highlights, CTA moments"),
    ("Ivory", "Base", f"#{IVORY}", "Backgrounds and premium fields"),
    ("Slate", "Secondary text", f"#{SLATE}", "Supporting copy and metadata"),
]
for token, role, hexv, use in palette:
    cells = color_table.add_row().cells
    for i, val in enumerate([token, role, hexv, use]):
        cells[i].text = val
        margins(cells[i])
        for p in cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                fmt_run(run, size=10, color="111827")

add_page_break()
section_title("04", "Typography", "Editorial, legible and product-friendly")
bullets([
    "Headline direction: elegant serif for premium brand moments.",
    "Body direction: clean sans serif for product, UI and explanations.",
    "Use the serif sparingly, mainly for the brand name and hero statements.",
    "Keep line lengths moderate and spacing calm.",
])

typography = doc.add_table(rows=1, cols=3)
typography.style = "Table Grid"
typography.autofit = False
for i, h in enumerate(["Use", "Recommended family", "Fallback"]):
    c = typography.rows[0].cells[i]
    c.text = h
    shade(c, "F3F4F6")
    margins(c)
for c in typography.rows[0].cells:
    for run in c.paragraphs[0].runs:
        fmt_run(run, size=10, bold=True, color=GRAFITO)
for row in [
    ("Brand / headlines", "Cormorant Garamond or Playfair Display", "Georgia"),
    ("Body / product", "Arial or Inter", "Segoe UI"),
    ("UI labels", "Arial Semibold or Inter Semibold", "Arial Bold"),
]:
    cells = typography.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val
        margins(cells[i])
        for p in cells[i].paragraphs:
            for run in p.runs:
                fmt_run(run, size=10, color="111827")

add_page_break()
section_title("05", "Applications", "How the system should appear in the real world")
bullets([
    "Web header: use the horizontal logo on light backgrounds.",
    "App icon: use the symbol only, centered with generous padding.",
    "Sales deck: keep the cover simple and premium with a lot of negative space.",
    "Social avatars: use the symbol, never the full wordmark.",
])

box = doc.add_table(rows=1, cols=1)
box.style = "Table Grid"
c = box.cell(0, 0)
c.text = "Brand voice: concise, polished and functional. The visual identity should make a restaurant feel modern without losing warmth."
shade(c, "F7F4EF")
margins(c, top=140, bottom=140)
for p in c.paragraphs:
    for run in p.runs:
        fmt_run(run, size=10.5, color=GRAFITO)

add_page_break()
section_title("06", "Do not", "Simple guardrails to protect the brand")
bullets([
    "Do not use saturated colors outside the palette.",
    "Do not put the logo inside a badge, circle or container unless the use case requires it.",
    "Do not mix font styles randomly.",
    "Do not add illustrations that compete with the logo.",
])

doc.save(OUT_PATH)
print(OUT_PATH)
